#!/usr/bin/env python3

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from openpyxl import load_workbook


BLUE = "2E5E78"
DARK_BLUE = "173B4F"
MUTED = "5D6870"
LIGHT_BLUE = "E8F0F4"
LIGHT_GRAY = "F2F4F5"
INK = "202529"
CAUTION = "FFF2CC"


def set_run_font(run, name="Times New Roman", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(element, fill):
    props = element.get_or_add_tcPr() if element.tag.endswith("tc") else element.get_or_add_pPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    props.append(node)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_field(paragraph, instruction, fallback=""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = fallback
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, text, separate, display, end])


def add_inline(paragraph, text, size=None, color=INK):
    token_pattern = re.compile(r"(\[PLACEHOLDER:[^\]]+\]|`[^`]+`|\*\*[^*]+\*\*|\[[^\]]+\]\([^)]+\))")
    for part in token_pattern.split(text):
        if not part:
            continue
        if part.startswith("[PLACEHOLDER:"):
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color="8A2D1F", bold=True)
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        elif part.startswith("`") and part.endswith("`"):
            inner = part[1:-1]
            if "[PLACEHOLDER:" in inner:
                add_inline(paragraph, inner, size=size, color=color)
            else:
                run = paragraph.add_run(inner)
                set_run_font(run, name="Consolas", size=(size or 11) - 0.5, color=DARK_BLUE)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        elif re.fullmatch(r"\[[^\]]+\]\([^)]+\)", part):
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part).groups()
            run = paragraph.add_run(f"{label} ({url})")
            set_run_font(run, size=size, color="0563C1")
            run.underline = True
        else:
            run = paragraph.add_run(part.replace("  ", " "))
            set_run_font(run, size=size, color=color)


def clean_markdown(text):
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    return text


def table_widths(column_count, total=9360):
    ratios = {
        2: [0.36, 0.64],
        3: [0.18, 0.38, 0.44],
        4: [0.18, 0.30, 0.22, 0.30],
        5: [0.13, 0.25, 0.18, 0.20, 0.24],
    }.get(column_count, [1 / column_count] * column_count)
    widths = [round(total * ratio) for ratio in ratios]
    widths[-1] += total - sum(widths)
    return widths


def add_markdown_table(doc, rows):
    headers = [clean_markdown(cell.strip()) for cell in rows[0]]
    body = rows[1:]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, table_widths(len(headers)))
    set_repeat_header(table.rows[0])
    set_cant_split(table.rows[0])
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade(cell._tc, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(paragraph, value, size=8.5, color=DARK_BLUE)
        for run in paragraph.runs:
            run.bold = True
    for source_row in body:
        row = table.add_row()
        set_cant_split(row)
        for index, value in enumerate(source_row):
            paragraph = row.cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(paragraph, clean_markdown(value.strip()), size=8.5)
    set_table_geometry(table, table_widths(len(headers)))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(doc, path, caption, width_inches=6.2):
    if not path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture = paragraph.add_run().add_picture(str(path), width=Inches(width_inches))
    picture._inline.docPr.set("title", caption)
    picture._inline.docPr.set("descr", caption)
    caption_p = doc.add_paragraph(style="Caption")
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(caption_p, caption, size=9, color=MUTED)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.5

    settings = {
        "Heading 1": (14, BLUE, 18, 10),
        "Heading 2": (14, BLUE, 12, 6),
        "Heading 3": (14, DARK_BLUE, 8, 4),
        "Heading 4": (14, DARK_BLUE, 6, 3),
    }
    for name, (size, color, before, after) in settings.items():
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.5

    code = doc.styles.add_style("Code block", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.font.size = Pt(8.5)
    code.paragraph_format.left_indent = Inches(0.2)
    code.paragraph_format.right_indent = Inches(0.2)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(3)

    caption = doc.styles["Caption"]
    caption.font.name = "Times New Roman"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_after = Pt(10)


def configure_section(section, landscape=False):
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)
    else:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.text = "KubeView final project report | Supervisor review copy"
    set_run_font(header.runs[0], size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Page ")
    set_run_font(run, size=8.5, color=MUTED)
    add_field(footer, " PAGE ", "1")


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(90)
    p.paragraph_format.space_after = Pt(18)
    add_inline(p, "KubeView", size=28, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(p, "Final project report", size=18, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    add_inline(p, "Submitted in partial fulfilment of the requirements for", size=11, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(p, "BSc Computer Science (Online Mode)", size=13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(p, "Birla Institute of Technology and Science, Pilani", size=12)

    for label, value in [
        ("Student", "Ankur Kalita | 2023EBCS782"),
        ("Student", "Pradyut Fogla | 2023EBCS788"),
        ("Student", "Varun Deep Saini | 2023EBCS663"),
        ("Supervisor", "[PLACEHOLDER: supervisor name, designation and organization]"),
        ("Semester", "[PLACEHOLDER: semester and academic year]"),
        ("Submission date", "[PLACEHOLDER: submission date]"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"{label}: ")
        set_run_font(run, size=11, color=MUTED, bold=True)
        add_inline(p, value, size=11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    shade(p._p, CAUTION)
    run = p.add_run("SUPERVISOR REVIEW COPY | PLACEHOLDERS REMAIN | SOURCE BASELINE 8ee601c")
    set_run_font(run, size=10, color="8A5A00", bold=True)
    doc.add_page_break()


def add_toc(doc):
    doc.add_heading("Table of contents", level=1)
    entries = [
        ("1. Introduction", "7"),
        ("2. Background and related systems", "10"),
        ("3. Requirements and success criteria", "12"),
        ("4. Architecture and design", "15"),
        ("5. Implementation", "20"),
        ("6. Security and privacy", "23"),
        ("7. Testing and validation", "25"),
        ("8. Results and evaluation", "27"),
        ("9. Project execution evidence and contribution record", "29"),
        ("10. Plagiarism, licensing, and attribution", "32"),
        ("11. Limitations and future work", "33"),
        ("12. Conclusion", "34"),
        ("References", "34"),
        ("Appendices A to D", "35"),
        ("Appendix E. Detailed test-case matrix", "38"),
        ("Appendix F. CI job evidence", "41"),
    ]
    for title, page in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        add_inline(p, f"{title}\t{page}", size=10.5)
    note = doc.add_paragraph()
    add_inline(note, "Review-copy pagination. Regenerate this table after the final metadata and evidence are inserted.", size=9, color=MUTED)
    doc.add_page_break()


def add_report_markdown(doc, markdown_path, screenshot_root):
    lines = markdown_path.read_text(encoding="utf8").splitlines()
    start = next(index for index, line in enumerate(lines) if line == "## Supervisor certificate")
    index = start
    in_code = False
    code_lines = []
    toc_added = False

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph(style="Code block")
                shade(p._p, LIGHT_GRAY)
                run = p.add_run("\n".join(code_lines))
                set_run_font(run, name="Consolas", size=8.5, color=INK)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if stripped.startswith("|"):
            raw_rows = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                    raw_rows.append(cells)
                index += 1
            if raw_rows:
                add_markdown_table(doc, raw_rows)
            continue

        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1)) - 1
            title = clean_markdown(heading.group(2))
            if title == "1. Introduction" and not toc_added:
                add_toc(doc)
                toc_added = True
            if (re.match(r"\d+\. ", title) and title != "1. Introduction") or title.startswith("Appendix D."):
                doc.add_page_break()
            if title == "4.9 Deployment design":
                doc.add_page_break()
            doc.add_heading(title, level=level)
            if title == "1.3 Unique value proposition":
                add_figure(doc, screenshot_root / "01-dashboard.png", "Figure 1. KubeView dashboard at the frozen source baseline")
            if title == "4.1 System structure":
                add_figure(
                    doc,
                    markdown_path.parent.parent / "architecture" / "system-context.png",
                    "Figure 2. KubeView system context and credential boundary",
                )
            if title == "5.3 Frontend state reconciliation":
                add_figure(doc, screenshot_root / "04-pod-detail.png", "Figure 3. Pod detail view")
                add_figure(doc, screenshot_root / "05-pod-logs.png", "Figure 4. Pod log view")
            index += 1
            continue

        if re.match(r"^- \[[ x]\] ", stripped):
            p = doc.add_paragraph(style="List Bullet")
            marker = "Complete: " if "[x]" in stripped[:6] else "Pending: "
            add_inline(p, "\u00a0" + marker + stripped[6:], size=10.5)
            index += 1
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, "\u00a0" + stripped[2:])
            index += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, "\u00a0" + numbered.group(1))
            index += 1
            continue
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.right_indent = Inches(0.2)
            shade(p._p, CAUTION)
            add_inline(p, stripped[2:], size=10)
            index += 1
            continue
        if not stripped:
            index += 1
            continue

        p = doc.add_paragraph()
        add_inline(p, stripped)
        index += 1


def add_test_appendix(doc, workbook_path):
    section = doc.add_section()
    configure_section(section, landscape=True)
    doc.add_heading("Appendix E. Detailed test-case matrix", level=1)
    p = doc.add_paragraph()
    add_inline(p, "Source: KubeView_Test_Case_Matrix.xlsx. Status applies to source commit 8ee601c.", size=9, color=MUTED)

    workbook = load_workbook(workbook_path, data_only=True)
    sheet = workbook["Test cases"]
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    widths = [800, 1450, 950, 4000, 1350, 1950]
    set_table_geometry(table, widths, indent=0)
    set_repeat_header(table.rows[0])
    set_cant_split(table.rows[0])
    headers = ["ID", "Requirement", "Level", "Scenario and expected result", "Automation", "Status, evidence and notes"]
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade(cell._tc, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(value)
        set_run_font(run, size=8, color=DARK_BLUE, bold=True)

    for values in sheet.iter_rows(min_row=2, values_only=True):
        test_id, requirement, level, scenario, expected, automation, status, evidence, notes = values
        row = table.add_row()
        set_cant_split(row)
        merged_values = [
            test_id,
            requirement,
            level,
            f"{scenario}\nExpected: {expected}",
            automation,
            f"{status}\nEvidence: {evidence or 'None recorded'}" + (f"\nNotes: {notes}" if notes else ""),
        ]
        for index, value in enumerate(merged_values):
            paragraph = row.cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            run = paragraph.add_run(str(value or ""))
            set_run_font(run, size=7.5, color=INK)
            if index in (0, 2, 4):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_geometry(table, widths, indent=0)

    doc.add_page_break()
    doc.add_heading("Appendix F. CI job evidence", level=1)
    ci_sheet = workbook["CI evidence"]
    ci_table = doc.add_table(rows=1, cols=4)
    ci_table.style = "Table Grid"
    ci_widths = [2300, 900, 2200, 5100]
    set_table_geometry(ci_table, ci_widths, indent=0)
    set_repeat_header(ci_table.rows[0])
    set_cant_split(ci_table.rows[0])
    headers = ["Check", "Result", "Execution window UTC", "Evidence URL"]
    for index, value in enumerate(headers):
        cell = ci_table.rows[0].cells[index]
        shade(cell._tc, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(value)
        set_run_font(run, size=8, color=DARK_BLUE, bold=True)
    for check, result, started, completed, url in ci_sheet.iter_rows(min_row=2, values_only=True):
        row = ci_table.add_row()
        set_cant_split(row)
        values = [check, result, f"{started} to {completed}", url]
        for index, value in enumerate(values):
            paragraph = row.cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(str(value or ""))
            set_run_font(run, size=7.5, color=INK)
            if index == 1:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_geometry(ci_table, ci_widths, indent=0)


def main():
    if len(sys.argv) != 5:
        raise SystemExit("usage: build_report_docx.py REPORT.md TESTS.xlsx SCREENSHOTS_DIR OUTPUT.docx")
    markdown_path = Path(sys.argv[1]).resolve()
    workbook_path = Path(sys.argv[2]).resolve()
    screenshot_root = Path(sys.argv[3]).resolve()
    output_path = Path(sys.argv[4]).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    doc.core_properties.title = "KubeView final project report"
    doc.core_properties.subject = "Supervisor review copy with placeholders"
    doc.core_properties.author = "Ankur Kalita; Pradyut Fogla; Varun Deep Saini"
    add_cover(doc)
    add_report_markdown(doc, markdown_path, screenshot_root)
    add_test_appendix(doc, workbook_path)

    settings = doc.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)
    doc.save(output_path)
    print(f"Wrote {output_path}")


if __name__ == "__main__":
    main()
